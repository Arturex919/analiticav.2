"""
ETL Module — Rental Analytics Platform
Handles multi-format ingestion, cleaning, and normalization.
"""

import re
import io
import pandas as pd
import numpy as np
from pathlib import Path
from typing import Optional


# ─────────────────────────────────────────────
# CONSTANTS
# ─────────────────────────────────────────────
STATUS_CONFIRMED  = ["Booked"]
STATUS_CANCELLED  = ["Declined", "Annuled", "annuled", "Anulled", "Cancelled"]
STATUS_BLOCKED    = ["Unavailable", "Cleaning"]
STATUS_OPEN       = ["Open", "Tentative"]

MONTH_ES = {
    1: "Enero", 2: "Febrero", 3: "Marzo", 4: "Abril",
    5: "Mayo", 6: "Junio", 7: "Julio", 8: "Agosto",
    9: "Septiembre", 10: "Octubre", 11: "Noviembre", 12: "Diciembre"
}
MONTH_ES_SHORT = {
    1:"Ene",2:"Feb",3:"Mar",4:"Abr",5:"May",6:"Jun",
    7:"Jul",8:"Ago",9:"Sep",10:"Oct",11:"Nov",12:"Dic"
}
DOW_ES = {
    "Monday":"Lunes","Tuesday":"Martes","Wednesday":"Miércoles",
    "Thursday":"Jueves","Friday":"Viernes","Saturday":"Sábado","Sunday":"Domingo"
}
DOW_ORDER = ["Lunes","Martes","Miércoles","Jueves","Viernes","Sábado","Domingo"]

SEASON_MAP = {
    12:"Invierno", 1:"Invierno", 2:"Invierno",
    3:"Primavera", 4:"Primavera", 5:"Primavera",
    6:"Verano", 7:"Verano", 8:"Verano",
    9:"Otoño", 10:"Otoño", 11:"Otoño"
}
QUARTER_MAP = {
    1:"Q1", 2:"Q1", 3:"Q1",
    4:"Q2", 5:"Q2", 6:"Q2",
    7:"Q3", 8:"Q3", 9:"Q3",
    10:"Q4", 11:"Q4", 12:"Q4"
}

# Spanish public holidays (month, day)
SPANISH_HOLIDAYS = {
    (1,1):"Año Nuevo",(1,6):"Reyes",(3,19):"San José",
    (4,18):"Viernes Santo",(5,1):"Día del Trabajo",
    (8,15):"Asunción",(10,12):"Fiesta Nacional",
    (11,1):"Todos los Santos",(12,6):"Constitución",
    (12,8):"Inmaculada",(12,25):"Navidad"
}

CHANNEL_ALIASES = {
    r"(?i)booking": "Booking.com",
    r"(?i)airbnb":  "Airbnb",
    r"(?i)homeaway|vrbo": "HomeAway/VRBO",
    r"(?i)^oh$":    "Other Holiday",
    r"(?i)manual|caty|maria": "Manual / Directo",
    r"(?i)limp":    "Limpieza/Bloqueo",
}


# ─────────────────────────────────────────────
# LOADERS
# ─────────────────────────────────────────────

def load_file(uploaded_file) -> pd.DataFrame:
    """Dispatch file loading by extension."""
    name = uploaded_file.name.lower()
    if name.endswith(".csv"):
        return _load_csv(uploaded_file)
    elif name.endswith((".xlsx", ".xls")):
        return _load_excel(uploaded_file)
    elif name.endswith(".pdf"):
        return _load_pdf(uploaded_file)
    elif name.endswith(".docx"):
        return _load_docx(uploaded_file)
    else:
        raise ValueError(f"Formato no soportado: {name}")


def _load_csv(f) -> pd.DataFrame:
    raw = f.read()
    # Intentamos primero con el separador detectado en rows.csv (punto y coma)
    for sep in [";", ",", "\t"]:
        try:
            df = pd.read_csv(io.BytesIO(raw), sep=sep, encoding="latin-1", quotechar='"', low_memory=False)
            if df.shape[1] > 2:
                # LIMPIEZA CRÍTICA: Eliminar triples comillas y espacios en cabeceras
                df.columns = [c.replace('"', '').strip() for c in df.columns]
                return df
        except Exception:
            pass
    return pd.read_csv(io.BytesIO(raw), sep=";", encoding="latin-1", quotechar='"')


def _load_excel(f) -> pd.DataFrame:
    try:
        return pd.read_excel(f, engine="openpyxl")
    except Exception:
        return pd.read_excel(f, engine="xlrd")


def _load_pdf(f) -> pd.DataFrame:
    """Extract tabular data from PDF using pypdf."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(f)
        text_blocks = []
        for page in reader.pages:
            text_blocks.append(page.extract_text())
        combined = "\n".join(text_blocks)
        # Attempt naive table parse: split on newlines and delimiters
        rows = [line.split() for line in combined.split("\n") if line.strip()]
        return pd.DataFrame(rows)
    except ImportError:
        raise ImportError("Instala pypdf: pip install pypdf")


def _load_docx(f) -> pd.DataFrame:
    """Extract tables from DOCX contracts/reports."""
    try:
        import docx
        doc = docx.Document(f)
        data = []
        for table in doc.tables:
            for row in table.rows:
                data.append([cell.text.strip() for cell in row.cells])
        if data:
            return pd.DataFrame(data[1:], columns=data[0])
        # Fall back to paragraphs
        rows = [{"texto": p.text} for p in doc.paragraphs if p.text.strip()]
        return pd.DataFrame(rows)
    except ImportError:
        raise ImportError("Instala python-docx: pip install python-docx")


# ─────────────────────────────────────────────
# COLUMN MAPPING
# ─────────────────────────────────────────────

REQUIRED_COLUMNS = {
    "id":        ["R_ID", "id", "booking_id", "reservation_id"],
    "status":    ["R_STATUS", "status", "estado"],
    "arrival":   ["R_ARRIVAL", "arrival", "llegada", "check_in"],
    "departure": ["R_DEPARTURE", "departure", "salida", "check_out"],
    "nights":    ["R_NIGHTS", "nights", "noches"],
    "amount":    ["R_TOTAL_AMOUNT", "amount", "total", "importe", "precio"],
    "property":  ["R_PROPERTY_NAME", "property", "propiedad", "alojamiento"],
    "source":    ["R_SOURCE", "source", "canal", "channel"],
    "guest":     ["R_GUEST_NAME", "guest", "huesped", "guest_name"],
    "owner":     ["NAME", "name", "propietario", "owner"],
    "created":   ["R_CREATED_AT", "created_at", "fecha_reserva"],
}


def _find_col(df: pd.DataFrame, candidates: list) -> Optional[str]:
    cols_lower = {c.lower(): c for c in df.columns}
    for c in candidates:
        if c in df.columns:
            return c
        if c.lower() in cols_lower:
            return cols_lower[c.lower()]
    return None


def map_columns(df: pd.DataFrame) -> dict:
    return {k: _find_col(df, v) for k, v in REQUIRED_COLUMNS.items()}


# ─────────────────────────────────────────────
# CLEANING & ENRICHMENT
# ─────────────────────────────────────────────

def clean_and_enrich(df: pd.DataFrame) -> pd.DataFrame:
    """Full ETL pipeline — returns a clean, enriched DataFrame."""
    df = df.copy()
    mapping = map_columns(df)

    # ── Rename to standard names ──
    rename = {v: k for k, v in mapping.items() if v and v in df.columns}
    df = df.rename(columns=rename)

    # ── Amount ──
    if "amount" in df.columns:
        # Limpieza ultra-robusta: eliminar todo lo que no sea dígito, coma o punto
        # Soporta formatos como € 1.234,56 o $1,234.56
        def _clean_curr(val):
            s = str(val).strip()
            if not s or s == "nan": return 0.0
            # Detectar si el punto es separador de miles (ej: 1.234,56)
            if "." in s and "," in s:
                if s.find(".") < s.find(","): # Estilo europeo: 1.234,56
                    s = s.replace(".", "").replace(",", ".")
                else: # Estilo USA: 1,234.56
                    s = s.replace(",", "")
            elif "," in s and "." not in s: # Estilo europeo simple: 1234,56
                s = s.replace(",", ".")
            
            s = re.sub(r'[^\d\.-]', '', s)
            try: return float(s)
            except: return 0.0

        df["amount"] = df["amount"].apply(_clean_curr)

    # ── Dates ──
    for col in ["arrival", "departure", "created"]:
        if col in df.columns:
            df[col] = _parse_dates(df[col])

    # ── Nights ──
    if "nights" in df.columns:
        df["nights"] = pd.to_numeric(df["nights"], errors="coerce")
        if "arrival" in df.columns and "departure" in df.columns:
            mask = df["nights"].isna()
            df.loc[mask, "nights"] = (
                (df.loc[mask, "departure"] - df.loc[mask, "arrival"]).dt.days
            )
    elif "arrival" in df.columns and "departure" in df.columns:
        df["nights"] = (df["departure"] - df["arrival"]).dt.days

    # ── Status normalisation ──
    if "status" in df.columns:
        df["status_group"] = df["status"].apply(_classify_status)

    # ── Channel normalisation ──
    if "source" in df.columns:
        df["channel"] = df["source"].apply(_normalize_channel)
    elif "source" not in df.columns:
        df["channel"] = "Desconocido"

    # ── Derived numeric fields ──
    if "amount" in df.columns and "nights" in df.columns:
        df["price_per_night"] = np.where(
            (df["nights"] > 0) & (df["amount"] > 0),
            df["amount"] / df["nights"],
            np.nan
        )

    # ── Temporal enrichment (based on arrival) ──
    if "arrival" in df.columns:
        dt = df["arrival"]
        df["month"]       = dt.dt.month
        df["month_name"]  = dt.dt.month.map(MONTH_ES)
        df["month_short"] = dt.dt.month.map(MONTH_ES_SHORT)
        df["year"]        = dt.dt.year
        df["quarter"]     = dt.dt.month.map(QUARTER_MAP)
        df["season"]      = dt.dt.month.map(SEASON_MAP)
        df["dow"]         = dt.dt.day_name().map(DOW_ES)
        df["week_num"]    = dt.dt.isocalendar().week.astype("Int64")
        df["is_weekend"]  = dt.dt.dayofweek >= 4   # Fri-Sun
        df["is_holiday"]  = dt.apply(_is_holiday)
        df["year_month"]  = dt.dt.to_period("M").astype(str)

    # ── Data Quality Flagging ──
    df["_is_outlier"] = _flag_outliers(df["amount"]) if "amount" in df else False
    df["_has_nulls"]  = df.isnull().any(axis=1)

    return df


def _flag_outliers(series: pd.Series) -> pd.Series:
    """Flags outliers using 3x IQR method."""
    if series.empty: return pd.Series(False, index=series.index)
    Q1, Q3 = series.quantile(0.25), series.quantile(0.75)
    IQR = Q3 - Q1
    return (series < Q1 - 3 * IQR) | (series > Q3 + 3 * IQR)


def impute_missing(df: pd.DataFrame, strategy: str = "mean") -> pd.DataFrame:
    """Imputes missing numeric values."""
    df = df.copy()
    num_cols = df.select_dtypes(include=[np.number]).columns
    for col in num_cols:
        if strategy == "mean":
            df[col] = df[col].fillna(df[col].mean())
        elif strategy == "median":
            df[col] = df[col].fillna(df[col].median())
        elif strategy == "zero":
            df[col] = df[col].fillna(0)
    return df


def _parse_dates(series: pd.Series) -> pd.Series:
    """Try multiple date formats."""
    formats = ["%d/%m/%Y", "%Y-%m-%d", "%d-%m-%Y", "%m/%d/%Y", "%d/%m/%Y %H:%M:%S"]
    for fmt in formats:
        try:
            parsed = pd.to_datetime(series, format=fmt, errors="coerce")
            if hasattr(parsed, 'notna') and parsed.notna().sum() / max(len(parsed), 1) > 0.5:
                return parsed
        except Exception:
            pass
    return pd.to_datetime(series, errors="coerce", dayfirst=True)


def _classify_status(s: str) -> str:
    if pd.isna(s):
        return "Desconocido"
    if s in STATUS_CONFIRMED:
        return "Confirmada"
    if s in STATUS_CANCELLED:
        return "Cancelada"
    if s in STATUS_BLOCKED:
        return "Bloqueada"
    if s in STATUS_OPEN:
        return "Abierta"
    return "Otro"


def _normalize_channel(s: str) -> str:
    if pd.isna(s) or str(s).strip() in ("", "nan", "<NULL>"):
        return "Directo / Sin canal"
    for pattern, label in CHANNEL_ALIASES.items():
        if re.search(pattern, str(s)):
            return label
    return str(s).strip().title()


def _is_holiday(dt) -> bool:
    if pd.isna(dt):
        return False
    return (dt.month, dt.day) in SPANISH_HOLIDAYS


# ─────────────────────────────────────────────
# FILTER HELPERS
# ─────────────────────────────────────────────

def filter_confirmed_paid(df: pd.DataFrame) -> pd.DataFrame:
    """Reservas confirmadas con ingreso real."""
    mask = (df.get("status_group") == "Confirmada") & (df.get("amount", 0) > 0)
    return df[mask].copy()


def quality_report(df_raw: pd.DataFrame, df_clean: pd.DataFrame) -> dict:
    return {
        "total_rows":      len(df_raw),
        "confirmed":       int((df_clean.get("status_group", pd.Series(dtype=str)) == "Confirmada").sum()),
        "with_revenue":    int((df_clean.get("amount", pd.Series(dtype=float)) > 0).sum()),
        "null_guest":      int(df_clean.get("guest", pd.Series(dtype=str)).isna().sum()),
        "null_source":     int(df_clean.get("source", pd.Series(dtype=str)).isna().sum()),
        "null_dates":      int(df_clean.get("arrival", pd.Series(dtype="datetime64[ns]")).isna().sum()),
        "date_range_min":  df_clean["arrival"].min() if "arrival" in df_clean else None,
        "date_range_max":  df_clean["arrival"].max() if "arrival" in df_clean else None,
        "properties":      df_clean["property"].nunique() if "property" in df_clean else 0,
        "channels":        df_clean["channel"].nunique() if "channel" in df_clean else 0,
    }
